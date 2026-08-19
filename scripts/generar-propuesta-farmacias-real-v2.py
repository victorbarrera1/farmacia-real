from __future__ import annotations

from datetime import datetime
import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REPO = Path(__file__).resolve().parents[1]
BASE_SCRIPT = REPO / "scripts" / "generar-propuesta-farmacias-real.py"
SPEC = importlib.util.spec_from_file_location("propuesta_base", BASE_SCRIPT)
u = importlib.util.module_from_spec(SPEC)
assert SPEC and SPEC.loader
SPEC.loader.exec_module(u)

OUT_DIR = REPO / "output" / "documentos"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_DOCX = OUT_DIR / "Propuesta_Plataforma_Digital_Farmacias_Real_Brida_v2.docx"

# Design preset: neighborhood_business_proposal.
# Header template: customer_pack.
# Named override: Brida editorial A4 — compact commercial proposal, 10 pages.
u.BRIDA_BLUE = "1769AA"
u.BRIDA_DARK = "10243A"
u.PLATFORM_NAVY = "172B4D"
u.PLATFORM_RED = "C72D37"
u.WA_GREEN = "147A50"
u.INK = "172033"
u.MUTED = "5E6B7A"
u.LIGHT_BLUE = "EAF3FA"
u.PALE = "F5F7F9"
u.PALE_GREEN = "EAF6F0"
u.PALE_RED = "FBEDEF"
u.LINE = "D8E1E8"
u.WHITE = "FFFFFF"

BLUE = u.BRIDA_BLUE
NAVY = u.BRIDA_DARK
INK = u.INK
MUTED = u.MUTED
PALE = u.PALE
LIGHT_BLUE = u.LIGHT_BLUE
PALE_GREEN = u.PALE_GREEN
PALE_RED = u.PALE_RED
RED = u.PLATFORM_RED
GREEN = u.WA_GREEN
LINE = u.LINE
WHITE = u.WHITE

TABLE_WIDTH = 9740
TABLE_INDENT = 0


def add_logo_natural(paragraph, width=2.56, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Place the supplied Brida wordmark without geometric distortion.

    The source is a 1000 px square JPEG. Its visible wordmark occupies roughly
    x=30..951 and y=352..595. This crop leaves 24 px of safety area around that
    content (970 x 292, aspect 3.322:1), and the output box uses the same ratio.
    """
    paragraph.alignment = align
    u.set_paragraph_spacing(paragraph, after=4, line=1.0)
    return u.add_picture_cropped(
        paragraph,
        u.BRIDA_LOGO,
        width,
        width / (970 / 292),
        {"l": 600, "r": 2400, "t": 32800, "b": 38000},
        "Logotipo oficial de Brida, conservando su proporción",
    )


def add_icon_natural(paragraph, width=0.22):
    # Square crop around the supplied icon; the glyph retains its native ratio.
    return u.add_picture_cropped(
        paragraph,
        u.BRIDA_ICON,
        width,
        width,
        {"l": 24500, "r": 24500, "t": 24500, "b": 24500},
        "Isotipo oficial de Brida",
    )


def configure_styles(doc):
    u.configure_styles(doc)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.15)
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing = 1.13

    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(22)
    h1.font.color.rgb = u.rgb(NAVY)
    h1.paragraph_format.space_after = Pt(7)

    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(12.4)
    h2.font.color.rgb = u.rgb(NAVY)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.size = Pt(9.25)
        style.paragraph_format.left_indent = Inches(0.21)
        style.paragraph_format.first_line_indent = Inches(-0.13)
        style.paragraph_format.space_after = Pt(2.7)
        style.paragraph_format.line_spacing = 1.08


def add_bottom_rule(paragraph, color=LINE, size=7, space=4):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def configure_section(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.2677)
    section.page_height = Inches(11.6929)
    section.top_margin = Inches(0.60)
    section.bottom_margin = Inches(0.56)
    section.left_margin = Inches(0.70)
    section.right_margin = Inches(0.70)
    section.header_distance = Inches(0.24)
    section.footer_distance = Inches(0.25)
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Inches(TABLE_WIDTH / 1440))
    u.set_table_geometry(table, [6600, 3140], indent=0)
    for cell in table.rows[0].cells:
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=0, bottom=45, start=0, end=0)
    left = u.clear_cell(table.cell(0, 0))
    r = left.add_run("FARMACIAS REAL  /  PROPUESTA DIGITAL")
    u.set_run_font(r, size=7.4, color=MUTED, bold=True)
    right = u.clear_cell(table.cell(0, 1))
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = right.add_run("AGOSTO 2026")
    u.set_run_font(r, size=7.4, color=BLUE, bold=True)
    add_bottom_rule(left, color=LINE, size=6, space=3)
    add_bottom_rule(right, color=LINE, size=6, space=3)

    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Inches(TABLE_WIDTH / 1440))
    u.set_table_geometry(table, [6850, 2890], indent=0)
    for cell in table.rows[0].cells:
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=20, bottom=0, start=0, end=0)
    left = u.clear_cell(table.cell(0, 0))
    add_icon_natural(left)
    r = left.add_run("  Brida · Propuesta comercial y operativa")
    u.set_run_font(r, size=7.8, color=MUTED)
    right = u.clear_cell(table.cell(0, 1))
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = right.add_run("Página ")
    u.set_run_font(r, size=7.8, color=MUTED)
    u.add_page_field(right, "PAGE")
    r = right.add_run(" de ")
    u.set_run_font(r, size=7.8, color=MUTED)
    u.add_page_field(right, "NUMPAGES")

    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    table = first_footer.add_table(rows=1, cols=2, width=Inches(TABLE_WIDTH / 1440))
    u.set_table_geometry(table, [6900, 2840], indent=0)
    for cell in table.rows[0].cells:
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=15, bottom=0, start=0, end=0)
    left = u.clear_cell(table.cell(0, 0))
    r = left.add_run("Propuesta preparada por Brida para Farmacias Real")
    u.set_run_font(r, size=7.8, color=MUTED)
    right = u.clear_cell(table.cell(0, 1))
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = right.add_run("18 de agosto de 2026")
    u.set_run_font(r, size=7.8, color=MUTED)


def section_header(doc, number, title, lead):
    table = doc.add_table(rows=1, cols=2)
    u.set_table_geometry(table, [820, 8920], indent=0)
    for cell in table.rows[0].cells:
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=0, bottom=25, start=0, end=70)
    p = u.clear_cell(table.cell(0, 0))
    r = p.add_run(number)
    u.set_run_font(r, size=10, color=BLUE, bold=True)
    p = u.clear_cell(table.cell(0, 1))
    u.set_paragraph_spacing(p, after=5, line=1.0, keep_next=True)
    r = p.add_run(title)
    u.set_run_font(r, size=22, color=NAVY, bold=True)
    p = doc.add_paragraph()
    u.set_paragraph_spacing(p, after=9, line=1.17)
    r = p.add_run(lead)
    u.set_run_font(r, size=10.8, color=MUTED)
    add_bottom_rule(p, color=LINE, size=6, space=5)


def bullet(container, text, size=9.2, color=INK, bold_prefix=None, after=2.7):
    p = container.add_paragraph(style="List Bullet")
    u.set_paragraph_spacing(p, after=after, line=1.08)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        u.set_run_font(r, size=size, color=NAVY, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        u.set_run_font(r, size=size, color=color)
    else:
        r = p.add_run(text)
        u.set_run_font(r, size=size, color=color)
    return p


def block(cell, title, lines, *, fill=WHITE, accent=BLUE, title_size=11.2,
          body_size=9.15, numbered=None, border=False):
    u.shade_cell(cell, fill)
    if border:
        u.set_cell_border(cell, color=LINE, size=7)
    else:
        u.remove_cell_borders(cell)
    u.set_cell_margins(cell, top=145, bottom=145, start=175, end=175)
    p = u.clear_cell(cell)
    u.set_paragraph_spacing(p, after=6, line=1.0, keep_next=True)
    if numbered:
        r = p.add_run(numbered + "  ")
        u.set_run_font(r, size=8.7, color=accent, bold=True)
    r = p.add_run(title)
    u.set_run_font(r, size=title_size, color=NAVY, bold=True)
    for line in lines:
        bullet(cell, line, size=body_size, color=MUTED)


def banner(doc, text, *, fill=NAVY, color=WHITE, accent=None, size=12.4, align=WD_ALIGN_PARAGRAPH.LEFT):
    table = doc.add_table(rows=1, cols=1)
    u.set_table_geometry(table, [TABLE_WIDTH], indent=0)
    cell = table.cell(0, 0)
    u.shade_cell(cell, fill)
    u.remove_cell_borders(cell)
    u.set_cell_margins(cell, top=175, bottom=175, start=220, end=220)
    p = u.clear_cell(cell)
    p.alignment = align
    u.set_paragraph_spacing(p, after=0, line=1.1)
    if accent:
        r = p.add_run(accent + "  ")
        u.set_run_font(r, size=size, color=BLUE, bold=True)
    r = p.add_run(text)
    u.set_run_font(r, size=size, color=color, bold=True)
    return table


def mini_stat(cell, label, value, *, accent=BLUE, fill=PALE):
    u.shade_cell(cell, fill)
    u.remove_cell_borders(cell)
    u.set_cell_margins(cell, top=115, bottom=115, start=135, end=135)
    p = u.clear_cell(cell)
    u.set_paragraph_spacing(p, after=3, line=1.0)
    r = p.add_run(label.upper())
    u.set_run_font(r, size=7.7, color=accent, bold=True)
    p = cell.add_paragraph()
    u.set_paragraph_spacing(p, after=0, line=1.05)
    r = p.add_run(value)
    u.set_run_font(r, size=9.5, color=NAVY, bold=True)


def page_cover(doc):
    top = doc.add_table(rows=1, cols=2)
    u.set_table_geometry(top, [5900, 3840], indent=0)
    for cell in top.rows[0].cells:
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
    p = u.clear_cell(top.cell(0, 0))
    add_logo_natural(p, width=2.56)
    p = u.clear_cell(top.cell(0, 1))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("PROPUESTA COMERCIAL\nY OPERATIVA")
    u.set_run_font(r, size=8.1, color=BLUE, bold=True)

    u.add_spacer(doc, 22)
    p = doc.add_paragraph()
    u.set_paragraph_spacing(p, after=12, line=0.98, keep_next=True)
    r = p.add_run("Propuesta de plataforma\ndigital para Farmacias Real")
    u.set_run_font(r, size=30, color=NAVY, bold=True)
    add_bottom_rule(p, color=BLUE, size=18, space=7)

    p = doc.add_paragraph()
    u.set_paragraph_spacing(p, after=17, line=1.18)
    r = p.add_run("Catálogo por sucursal, reserva por WhatsApp y control operativo en una sola experiencia.")
    u.set_run_font(r, size=14.2, color=MUTED)

    statement = doc.add_table(rows=1, cols=2)
    u.set_table_geometry(statement, [590, 9150], indent=0)
    left, right = statement.cell(0, 0), statement.cell(0, 1)
    u.shade_cell(left, BLUE)
    u.remove_cell_borders(left)
    u.set_cell_margins(left, top=0, bottom=0, start=0, end=0)
    u.shade_cell(right, LIGHT_BLUE)
    u.remove_cell_borders(right)
    u.set_cell_margins(right, top=210, bottom=210, start=230, end=230)
    p = u.clear_cell(right)
    u.set_paragraph_spacing(p, after=0, line=1.13)
    r = p.add_run("Una farmacia de barrio con la facilidad digital que hoy esperan sus vecinos.")
    u.set_run_font(r, size=15, color=NAVY, bold=True)

    u.add_spacer(doc, 18)
    u.add_kicker(doc, "Una propuesta integrada", color=BLUE, after=7)
    features = doc.add_table(rows=1, cols=3)
    u.set_table_geometry(features, [3246, 3247, 3247], indent=0)
    items = (
        ("01", "Catálogo por local", "Búsqueda, filtros, stock y precios por sucursal."),
        ("02", "Reserva conversada", "Pedido preparado y coordinación directa por WhatsApp."),
        ("03", "Control operativo", "Panel para administración general y encargados."),
    )
    for idx, (num, title, body) in enumerate(items):
        block(features.cell(0, idx), title, [body], fill=PALE if idx != 1 else WHITE,
              accent=GREEN if idx == 1 else BLUE, numbered=num, body_size=8.8, border=True)

    u.add_spacer(doc, 24)
    meta = doc.add_table(rows=2, cols=2)
    u.set_table_geometry(meta, [4870, 4870], indent=0)
    rows = (
        ("PREPARADO PARA", "Farmacias Real", "PREPARADO POR", "Brida"),
        ("COBERTURA", "4 sucursales · Independencia y Ñuñoa", "FECHA", "18 de agosto de 2026"),
    )
    for ridx, row in enumerate(rows):
        for cidx in range(2):
            cell = meta.cell(ridx, cidx)
            u.remove_cell_borders(cell)
            u.set_cell_margins(cell, top=100, bottom=100, start=0 if cidx == 0 else 170, end=120)
            p = u.clear_cell(cell)
            r = p.add_run(row[cidx * 2] + "\n")
            u.set_run_font(r, size=7.5, color=BLUE, bold=True)
            r = p.add_run(row[cidx * 2 + 1])
            u.set_run_font(r, size=9.3, color=NAVY, bold=True)
            add_bottom_rule(p, color=LINE, size=5, space=2)

    u.add_spacer(doc, 20)
    u.add_kicker(doc, "La propuesta recorre", color=BLUE, after=6)
    roadmap = doc.add_table(rows=1, cols=4)
    u.set_table_geometry(roadmap, [2435] * 4, indent=0)
    for idx, (label, text) in enumerate((
        ("01 · EXPERIENCIA", "Cómo encuentra y coordina el vecino"),
        ("02 · OPERACIÓN", "Cómo trabaja cada sucursal"),
        ("03 · ACTIVACIÓN", "Cómo llevarla al barrio"),
        ("04 · PUESTA EN MARCHA", "Qué falta para publicar"),
    )):
        cell = roadmap.cell(0, idx)
        u.shade_cell(cell, NAVY if idx % 2 == 0 else BLUE)
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=120, bottom=120, start=90, end=90)
        p = u.clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label + "\n")
        u.set_run_font(r, size=7.2, color=WHITE, bold=True)
        r = p.add_run(text)
        u.set_run_font(r, size=7.6, color="EAF1F5")


def page_executive(doc):
    section_header(
        doc,
        "01",
        "La decisión comienza antes de la visita",
        "Antes de salir, el vecino quiere saber qué necesita, en qué local está disponible y cómo coordinar. La plataforma reúne vitrina pública, WhatsApp y gestión por sucursal.",
    )

    frame = doc.add_table(rows=1, cols=1)
    u.set_table_geometry(frame, [TABLE_WIDTH], indent=0)
    cell = frame.cell(0, 0)
    u.shade_cell(cell, PALE)
    u.remove_cell_borders(cell)
    u.set_cell_margins(cell, top=80, bottom=75, start=80, end=80)
    p = u.clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(u.FARMACIA_OG), width=Inches(6.50))
    shape._inline.docPr.set("descr", "Visual real del prototipo de Farmacias Real")
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    u.set_paragraph_spacing(p, before=2, after=0, line=1.0)
    r = p.add_run("Visual del prototipo funcional · La identidad de Farmacias Real mostrada aquí es provisional")
    u.set_run_font(r, size=7.4, color=MUTED, italic=True)

    u.add_spacer(doc, 9)
    outcomes = doc.add_table(rows=1, cols=3)
    u.set_table_geometry(outcomes, [3246, 3247, 3247], indent=0)
    items = (
        ("VECINO", "Encuentra y coordina", "Consulta antes de trasladarse y elige dónde retirar."),
        ("EQUIPO", "Recibe más orden", "El mensaje incluye productos, cantidades y local."),
        ("DUEÑO", "Controla cuatro locales", "Administra catálogo, stock y precio desde una base."),
    )
    for idx, (label, title, body) in enumerate(items):
        cell = outcomes.cell(0, idx)
        u.shade_cell(cell, WHITE)
        u.set_cell_border(cell, color=LINE, size=6, sides=("bottom",))
        u.set_cell_margins(cell, top=120, bottom=130, start=120, end=130)
        p = u.clear_cell(cell)
        r = p.add_run(label)
        u.set_run_font(r, size=7.7, color=BLUE if idx != 1 else GREEN, bold=True)
        p = cell.add_paragraph()
        u.set_paragraph_spacing(p, after=4, line=1.0)
        r = p.add_run(title)
        u.set_run_font(r, size=10.7, color=NAVY, bold=True)
        p = cell.add_paragraph()
        u.set_paragraph_spacing(p, after=0, line=1.1)
        r = p.add_run(body)
        u.set_run_font(r, size=8.7, color=MUTED)

    u.add_spacer(doc, 9)
    banner(doc, "No es un e-commerce: la web prepara la consulta y la farmacia confirma, cobra y entrega en el local.", fill=NAVY, size=11.8)
    u.add_spacer(doc, 8)
    branches = doc.add_table(rows=1, cols=4)
    u.set_table_geometry(branches, [2435] * 4, indent=0)
    for idx, (name, comuna) in enumerate((
        ("Independencia 1443", "Independencia"),
        ("Sevilla 1201", "Independencia"),
        ("Santa María 1789", "Independencia"),
        ("Simón Bolívar 3751", "Ñuñoa"),
    )):
        mini_stat(branches.cell(0, idx), comuna, name, fill=LIGHT_BLUE if idx < 3 else PALE)

    u.add_spacer(doc, 8)
    u.add_subheading(doc, "Una ventaja equilibrada", size=11.5, after=5)
    compare = doc.add_table(rows=1, cols=3)
    u.set_table_geometry(compare, [3246, 3247, 3247], indent=0)
    comparisons = (
        ("GRAN CADENA", "Facilidad digital, con una relación más distribuida."),
        ("FARMACIA LOCAL SIN PLATAFORMA", "Cercanía, pero con más consultas manuales y poca información previa."),
        ("FARMACIAS REAL CON PLATAFORMA", "Facilidad digital por local sin perder la conversación directa."),
    )
    for idx, (title, body) in enumerate(comparisons):
        cell = compare.cell(0, idx)
        u.shade_cell(cell, PALE_GREEN if idx == 2 else PALE)
        u.set_cell_border(cell, color=LINE, size=5, sides=("bottom",))
        u.set_cell_margins(cell, top=95, bottom=95, start=105, end=105)
        p = u.clear_cell(cell)
        r = p.add_run(title + "\n")
        u.set_run_font(r, size=7.2, color=GREEN if idx == 2 else BLUE, bold=True)
        r = p.add_run(body)
        u.set_run_font(r, size=8.05, color=MUTED)


def page_journey(doc):
    section_header(
        doc,
        "02",
        "Del buscador al mesón, sin pasos extra",
        "El cliente recorre una secuencia conocida y breve desde cualquier teléfono, tablet o computador. La farmacia conserva la confirmación humana.",
    )

    flow = doc.add_table(rows=2, cols=3)
    u.set_table_geometry(flow, [3246, 3247, 3247], indent=0)
    steps = (
        ("01", "Elige sucursal", "Define el local donde desea retirar."),
        ("02", "Busca o navega", "Usa nombre, marca, principio activo o categorías."),
        ("03", "Afina resultados", "Combina disponibilidad, bioequivalencia, receta y precio."),
        ("04", "Revisa el producto", "Ve presentación, laboratorio, stock y precio referencial."),
        ("05", "Prepara el pedido", "Agrega productos y ajusta cantidades."),
        ("06", "Coordina y retira", "WhatsApp confirma; el pago y la entrega son presenciales."),
    )
    for idx, (num, title, body) in enumerate(steps):
        block(flow.cell(idx // 3, idx % 3), title, [body],
              fill=LIGHT_BLUE if idx < 3 else WHITE,
              accent=GREEN if idx == 5 else BLUE,
              numbered=num, body_size=8.8, border=True)

    u.add_spacer(doc, 9)
    banner(doc, "Sin cuenta, contraseña, datos bancarios, formularios extensos ni aplicación que instalar.", fill=BLUE, size=11.7)

    u.add_spacer(doc, 9)
    u.add_subheading(doc, "Buen uso en situaciones reales", size=12.1, after=6)
    uses = doc.add_table(rows=2, cols=2)
    u.set_table_geometry(uses, [4870, 4870], indent=0)
    items = (
        ("Antes de trasladarse", "Un vecino revisa stock y evita visitar un local sin disponibilidad."),
        ("Desde el hogar", "Un cuidador prepara una solicitud para un familiar con nombres y cantidades correctos."),
        ("Otra sucursal", "Si falta un producto, el cliente revisa en qué otro local aparece disponible."),
        ("Atención más clara", "El personal recibe el pedido, el local de retiro y el total referencial ya ordenados."),
    )
    for idx, (title, body) in enumerate(items):
        cell = uses.cell(idx // 2, idx % 2)
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=105, bottom=105, start=0 if idx % 2 == 0 else 175, end=120)
        p = u.clear_cell(cell)
        r = p.add_run(title + ". ")
        u.set_run_font(r, size=9.2, color=NAVY, bold=True)
        r = p.add_run(body)
        u.set_run_font(r, size=9.2, color=MUTED)
        add_bottom_rule(p, color=LINE, size=5, space=2)

    u.add_spacer(doc, 7)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    u.set_paragraph_spacing(p, after=0, line=1.1)
    r = p.add_run("Diseño directo: textos legibles, objetivos táctiles, navegación por teclado y sin movimientos inesperados.")
    u.set_run_font(r, size=8.7, color=MUTED, italic=True)

    u.add_spacer(doc, 13)
    principles = doc.add_table(rows=1, cols=3)
    u.set_table_geometry(principles, [3246, 3247, 3247], indent=0)
    for idx, (title, body) in enumerate((
        ("MÓVIL PRIMERO", "La búsqueda, la sucursal y el pedido permanecen al alcance."),
        ("CLARIDAD PARA MÁS PERSONAS", "Textos legibles, botones directos y navegación compatible con teclado."),
        ("SIN FRICCIÓN DIGITAL", "Las funciones importantes se conservan sin instalar una aplicación."),
    )):
        cell = principles.cell(0, idx)
        u.shade_cell(cell, LIGHT_BLUE if idx == 0 else (PALE_GREEN if idx == 2 else PALE))
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=120, bottom=120, start=125, end=125)
        p = u.clear_cell(cell)
        r = p.add_run(title + "\n")
        u.set_run_font(r, size=7.6, color=GREEN if idx == 2 else BLUE, bold=True)
        r = p.add_run(body)
        u.set_run_font(r, size=8.4, color=MUTED)


def page_public_experience(doc):
    section_header(
        doc,
        "03",
        "Todo lo que contiene la experiencia pública",
        "Las funciones están agrupadas para que encontrar, comparar y coordinar sea rápido. Cada elemento se relaciona con una decisión concreta del cliente.",
    )

    grid = doc.add_table(rows=2, cols=2)
    u.set_table_geometry(grid, [4870, 4870], indent=0)
    blocks = (
        ("Encontrar", [
            "Búsqueda por nombre, marca o laboratorio, principio activo y presentación.",
            "Sugerencias mientras se escribe y resultados en la misma página.",
            "Navegación por siete categorías de productos.",
        ], LIGHT_BLUE, BLUE),
        ("Afinar", [
            "Filtros por disponibilidad, bioequivalencia, condición de venta, precio y laboratorio.",
            "Orden por recomendados, menor o mayor precio y orden alfabético.",
            "Indicadores de filtros activos y limpieza con un toque.",
        ], WHITE, BLUE),
        ("Decidir", [
            "Ficha con presentación, laboratorio, principio activo y descripción.",
            "Precio referencial, disponibilidad y sellos informativos.",
            "Aviso de stock visible en otros locales cuando corresponde.",
        ], WHITE, RED),
        ("Coordinar", [
            "Pedido con cantidades, total referencial, sucursal y dirección.",
            "Consulta puntual o pedido completo dirigido al WhatsApp del local.",
            "Retiro, pago y validaciones permanecen en la sucursal.",
        ], PALE_GREEN, GREEN),
    )
    for idx, (title, lines, fill, accent) in enumerate(blocks):
        block(grid.cell(idx // 2, idx % 2), title, lines, fill=fill, accent=accent,
              numbered=f"0{idx + 1}", body_size=8.85, border=True)

    u.add_spacer(doc, 10)
    u.add_subheading(doc, "Categorías implementadas", size=12.1, after=6)
    cats = doc.add_table(rows=1, cols=7)
    u.set_table_geometry(cats, [1391, 1391, 1391, 1391, 1392, 1392, 1392], indent=0)
    labels = ("Medicamentos", "Dermocosmética", "Perfumería", "Vitaminas", "Mamá y bebé", "Cuidado personal", "Equipos y control")
    for idx, label in enumerate(labels):
        cell = cats.cell(0, idx)
        u.shade_cell(cell, NAVY if idx % 2 == 0 else BLUE)
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=105, bottom=105, start=45, end=45)
        p = u.clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        u.set_run_font(r, size=7.35, color=WHITE, bold=True)

    u.add_spacer(doc, 8)
    p = doc.add_paragraph()
    u.set_paragraph_spacing(p, after=7, line=1.12)
    r = p.add_run("Perfumería amplía el uso del catálogo")
    u.set_run_font(r, size=9.5, color=NAVY, bold=True)
    r = p.add_run(": incluye fragancias, colonias, desodorantes, cuidado corporal y sets de regalo dentro del catálogo demostrativo.")
    u.set_run_font(r, size=9.5, color=MUTED)

    note = doc.add_table(rows=1, cols=2)
    u.set_table_geometry(note, [2180, 7560], indent=0)
    left, right = note.cell(0, 0), note.cell(0, 1)
    u.shade_cell(left, RED)
    u.remove_cell_borders(left)
    u.set_cell_margins(left, top=145, bottom=145, start=140, end=140)
    p = u.clear_cell(left)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BIOEQUIVALENCIA")
    u.set_run_font(r, size=8.7, color=WHITE, bold=True)
    u.shade_cell(right, PALE_RED)
    u.remove_cell_borders(right)
    u.set_cell_margins(right, top=120, bottom=120, start=180, end=180)
    p = u.clear_cell(right)
    r = p.add_run("El filtro muestra la clasificación registrada en el catálogo. No recomienda sustituciones ni reemplaza la orientación profesional.")
    u.set_run_font(r, size=9, color=INK, bold=True)

    u.add_spacer(doc, 10)
    u.add_subheading(doc, "Información visible para decidir", size=11.8, after=5)
    visible = doc.add_table(rows=2, cols=3)
    u.set_table_geometry(visible, [3246, 3247, 3247], indent=0)
    for idx, (label, value) in enumerate((
        ("PRODUCTO", "Nombre y presentación"),
        ("ORIGEN", "Laboratorio o marca"),
        ("COMPOSICIÓN", "Principio activo registrado"),
        ("CONDICIÓN", "Receta y otros sellos"),
        ("SUCURSAL", "Stock y alternativas"),
        ("VALOR", "Precio referencial publicado"),
    )):
        mini_stat(visible.cell(idx // 3, idx % 3), label, value,
                  accent=RED if idx == 3 else BLUE,
                  fill=PALE_RED if idx == 3 else (LIGHT_BLUE if idx < 3 else PALE))


def page_branches(doc):
    section_header(
        doc,
        "04",
        "Cuando cambia el local, cambia la información que importa",
        "La misma base de productos puede reflejar diferencias reales por sucursal. El cliente ve el contexto del local donde desea retirar.",
    )

    data = doc.add_table(rows=1, cols=6)
    u.set_table_geometry(data, [1623, 1623, 1623, 1623, 1624, 1624], indent=0)
    for idx, label in enumerate(("Catálogo visible", "Disponibilidad", "Precio local", "Dirección", "Horario", "WhatsApp")):
        cell = data.cell(0, idx)
        u.shade_cell(cell, NAVY if idx % 2 == 0 else BLUE)
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=125, bottom=125, start=55, end=55)
        p = u.clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        u.set_run_font(r, size=7.9, color=WHITE, bold=True)

    u.add_spacer(doc, 11)
    table = doc.add_table(rows=5, cols=4)
    u.set_table_geometry(table, [2350, 2180, 3320, 1890], indent=0)
    headers = ("Sucursal", "Comuna", "Dirección", "Contacto")
    for idx, label in enumerate(headers):
        cell = table.cell(0, idx)
        u.shade_cell(cell, LIGHT_BLUE)
        u.set_cell_border(cell, color=LINE, size=6, sides=("bottom",))
        u.set_cell_margins(cell, top=105, bottom=105, start=100, end=100)
        p = u.clear_cell(cell)
        r = p.add_run(label.upper())
        u.set_run_font(r, size=7.6, color=BLUE, bold=True)
    rows = (
        ("Independencia 1443", "Independencia", "Av. Independencia 1443", "WhatsApp definido"),
        ("Sevilla 1201", "Independencia", "Calle Sevilla 1201–1205", "WhatsApp definido"),
        ("Santa María 1789", "Independencia", "Av. Domingo Santa María 1789", "Provisional"),
        ("Simón Bolívar 3751", "Ñuñoa", "Av. Simón Bolívar 3751-A", "Provisional"),
    )
    for ridx, row in enumerate(rows, start=1):
        for cidx, value in enumerate(row):
            cell = table.cell(ridx, cidx)
            u.shade_cell(cell, WHITE if ridx % 2 else PALE)
            u.set_cell_border(cell, color=LINE, size=5, sides=("bottom",))
            u.set_cell_margins(cell, top=105, bottom=105, start=100, end=100)
            p = u.clear_cell(cell)
            r = p.add_run(value)
            u.set_run_font(r, size=8.45, color=RED if value == "Provisional" else INK,
                           bold=(cidx == 0 or value == "Provisional"))

    u.add_spacer(doc, 11)
    benefits = doc.add_table(rows=1, cols=2)
    u.set_table_geometry(benefits, [4870, 4870], indent=0)
    block(benefits.cell(0, 0), "Para el cliente", [
        "Evita ir al local equivocado o trasladarse sin certeza de disponibilidad.",
        "Puede revisar otra sucursal cuando el producto aparece sin stock.",
        "Recibe dirección y canal de contacto del retiro elegido.",
    ], fill=LIGHT_BLUE, accent=BLUE, body_size=9.0)
    block(benefits.cell(0, 1), "Para la farmacia", [
        "Publica diferencias reales sin duplicar todo el catálogo.",
        "Orienta la demanda hacia el local con disponibilidad.",
        "Ajusta precio y visibilidad de cada sucursal de forma independiente.",
    ], fill=PALE_GREEN, accent=GREEN, body_size=9.0)

    u.add_spacer(doc, 9)
    banner(doc, "Antes de publicar: confirmar los WhatsApp definitivos de Santa María 1789 y Simón Bolívar 3751.", fill=PALE_RED, color=RED, size=10.2)

    u.add_spacer(doc, 10)
    u.add_subheading(doc, "Buen uso de la información por local", size=11.8, after=5)
    cadence = doc.add_table(rows=1, cols=3)
    u.set_table_geometry(cadence, [3246, 3247, 3247], indent=0)
    for idx, (title, body) in enumerate((
        ("PUBLICAR", "Mostrar solo los productos que el local realmente maneja."),
        ("ACTUALIZAR", "Revisar unidades, precio y horario con una rutina definida."),
        ("ORIENTAR", "Usar otra sucursal como alternativa cuando existe disponibilidad."),
    )):
        cell = cadence.cell(0, idx)
        u.shade_cell(cell, LIGHT_BLUE if idx == 0 else (PALE_GREEN if idx == 2 else PALE))
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=120, bottom=120, start=125, end=125)
        p = u.clear_cell(cell)
        r = p.add_run(title + "\n")
        u.set_run_font(r, size=7.6, color=GREEN if idx == 2 else BLUE, bold=True)
        r = p.add_run(body)
        u.set_run_font(r, size=8.4, color=MUTED)


def page_whatsapp(doc):
    section_header(
        doc,
        "05",
        "La tecnología prepara el pedido; la farmacia mantiene la relación",
        "WhatsApp conecta el catálogo con una conversación humana. La solicitud llega ordenada, pero la reserva solo existe cuando la farmacia la confirma.",
    )

    flow = doc.add_table(rows=1, cols=7)
    u.set_table_geometry(flow, [2100, 310, 2100, 310, 2100, 310, 2510], indent=0)
    nodes = (
        (0, "CATÁLOGO", "Productos y cantidades", NAVY),
        (2, "MENSAJE", "Datos preparados", BLUE),
        (4, "CONFIRMACIÓN", "Respuesta humana", GREEN),
        (6, "LOCAL", "Retiro y pago", NAVY),
    )
    for idx, title, body, fill in nodes:
        cell = flow.cell(0, idx)
        u.shade_cell(cell, fill)
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=135, bottom=135, start=90, end=90)
        p = u.clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title + "\n")
        u.set_run_font(r, size=8.1, color=WHITE, bold=True)
        r = p.add_run(body)
        u.set_run_font(r, size=7.7, color="EAF1F5")
    for idx in (1, 3, 5):
        cell = flow.cell(0, idx)
        u.remove_cell_borders(cell)
        p = u.clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("›")
        u.set_run_font(r, size=17, color=BLUE, bold=True)

    u.add_spacer(doc, 10)
    info = doc.add_table(rows=1, cols=2)
    u.set_table_geometry(info, [4870, 4870], indent=0)
    block(info.cell(0, 0), "El mensaje incluye", [
        "Productos seleccionados y cantidades.",
        "Total referencial del pedido.",
        "Sucursal y dirección de retiro.",
        "Aviso de pago y entrega presenciales.",
    ], fill=LIGHT_BLUE, accent=BLUE, body_size=9.0)
    block(info.cell(0, 1), "El buen uso del canal", [
        "El cliente revisa antes de enviar.",
        "El equipo confirma stock, condición de venta y valor final.",
        "La conversación aclara dudas sin reescribir el pedido.",
        "La visita se coordina antes del traslado.",
    ], fill=PALE_GREEN, accent=GREEN, body_size=9.0)

    u.add_spacer(doc, 10)
    guards = doc.add_table(rows=1, cols=3)
    u.set_table_geometry(guards, [3246, 3247, 3247], indent=0)
    for idx, (title, body) in enumerate((
        ("SIN PAGO ONLINE", "El valor final se confirma y se paga en caja."),
        ("SIN DESPACHO", "La entrega se realiza en la sucursal elegida."),
        ("SIN CONFIRMACIÓN AUTOMÁTICA", "Enviar el mensaje no garantiza una reserva."),
    )):
        block(guards.cell(0, idx), title, [body], fill=PALE_RED if idx == 2 else PALE,
              accent=RED if idx == 2 else BLUE, body_size=8.7, border=True)

    u.add_spacer(doc, 10)
    prescription = doc.add_table(rows=1, cols=2)
    u.set_table_geometry(prescription, [2380, 7360], indent=0)
    left, right = prescription.cell(0, 0), prescription.cell(0, 1)
    u.shade_cell(left, RED)
    u.remove_cell_borders(left)
    u.set_cell_margins(left, top=155, bottom=155, start=140, end=140)
    p = u.clear_cell(left)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RESERVAR CON RECETA")
    u.set_run_font(r, size=8.8, color=WHITE, bold=True)
    u.shade_cell(right, PALE_RED)
    u.remove_cell_borders(right)
    u.set_cell_margins(right, top=125, bottom=125, start=180, end=180)
    p = u.clear_cell(right)
    r = p.add_run("Los productos sujetos a receta se presentan de manera neutral. La receta y la entrega deben validarse presencialmente por el profesional correspondiente.")
    u.set_run_font(r, size=8.9, color=INK, bold=True)

    u.add_spacer(doc, 8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    u.set_paragraph_spacing(p, after=0, line=1.08)
    r = p.add_run("Esta propuesta no constituye asesoría legal ni garantiza por sí sola cumplimiento regulatorio; las decisiones definitivas deben ser revisadas por la farmacia y sus asesores cuando corresponda.")
    u.set_run_font(r, size=7.9, color=MUTED, italic=True)

    u.add_spacer(doc, 11)
    u.add_subheading(doc, "Protocolo recomendado para cada solicitud", size=11.8, after=5)
    protocol = doc.add_table(rows=1, cols=3)
    u.set_table_geometry(protocol, [3246, 3247, 3247], indent=0)
    for idx, (title, body) in enumerate((
        ("1 · REVISAR", "Confirmar productos, cantidades y sucursal elegida."),
        ("2 · RESPONDER", "Informar disponibilidad, condición de venta y valor final."),
        ("3 · COORDINAR", "Indicar retiro, horario y documentación que corresponda."),
    )):
        cell = protocol.cell(0, idx)
        u.shade_cell(cell, LIGHT_BLUE if idx == 0 else (PALE_GREEN if idx == 2 else PALE))
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=110, bottom=110, start=120, end=120)
        p = u.clear_cell(cell)
        r = p.add_run(title + "\n")
        u.set_run_font(r, size=7.6, color=GREEN if idx == 2 else BLUE, bold=True)
        r = p.add_run(body)
        u.set_run_font(r, size=8.3, color=MUTED)


def page_panel(doc):
    section_header(
        doc,
        "06",
        "Control general para el dueño; autonomía acotada para cada local",
        "El panel separa las tareas de administración general de las actualizaciones cotidianas de cada sucursal. Los permisos ayudan a reducir cambios accidentales.",
    )

    table = doc.add_table(rows=7, cols=3)
    u.set_table_geometry(table, [2440, 3650, 3650], indent=0)
    headers = ("Ámbito", "Administración general", "Encargado de sucursal")
    for idx, label in enumerate(headers):
        cell = table.cell(0, idx)
        u.shade_cell(cell, NAVY if idx == 0 else BLUE)
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=115, bottom=115, start=100, end=100)
        p = u.clear_cell(cell)
        r = p.add_run(label.upper())
        u.set_run_font(r, size=7.6, color=WHITE, bold=True)
    rows = (
        ("Resumen", "Ve la operación consolidada y por local.", "Ve los indicadores de su sucursal."),
        ("Productos", "Crea, edita y elimina; asigna categoría, bioequivalencia y receta.", "Consulta el catálogo; no crea ni elimina productos."),
        ("Stock y precio", "Actualiza cualquier sucursal y define visibilidad local.", "Actualiza stock, precio y visibilidad de su local."),
        ("Sucursales", "Edita datos, horarios, teléfonos y WhatsApp.", "No modifica otros locales ni su configuración general."),
        ("Solicitudes", "Revisa el historial de todas las sucursales.", "Revisa únicamente las solicitudes de su local."),
        ("Datos y respaldo", "Importa y exporta CSV; descarga y restaura respaldos.", "Sin acceso a operaciones globales de respaldo."),
    )
    for ridx, row in enumerate(rows, start=1):
        for cidx, value in enumerate(row):
            cell = table.cell(ridx, cidx)
            u.shade_cell(cell, WHITE if ridx % 2 else PALE)
            u.set_cell_border(cell, color=LINE, size=5, sides=("bottom",))
            u.set_cell_margins(cell, top=95, bottom=95, start=105, end=105)
            p = u.clear_cell(cell)
            r = p.add_run(value)
            u.set_run_font(r, size=8.15, color=NAVY if cidx == 0 else MUTED, bold=(cidx == 0))

    u.add_spacer(doc, 10)
    u.add_subheading(doc, "Buen uso operativo", size=12.1, after=6)
    good_use = doc.add_table(rows=1, cols=2)
    u.set_table_geometry(good_use, [4870, 4870], indent=0)
    block(good_use.cell(0, 0), "Dueño o administración", [
        "Revisar quiebres, stock bajo y diferencias entre locales.",
        "Hacer cambios masivos e importaciones con respaldo previo.",
        "Mantener datos de producto con información verificada.",
    ], fill=LIGHT_BLUE, accent=BLUE, body_size=8.85)
    block(good_use.cell(0, 1), "Encargado de sucursal", [
        "Actualizar unidades, precio y visibilidad al inicio o cierre de turno.",
        "Revisar las solicitudes asociadas a su local.",
        "Informar inconsistencias al administrador general.",
    ], fill=PALE_GREEN, accent=GREEN, body_size=8.85)

    u.add_spacer(doc, 9)
    banner(doc, "Administrar la página debe sentirse tan directo como actualizar una planilla, pero con controles que ayudan a evitar errores.", fill=BLUE, size=11.2)
    u.add_spacer(doc, 7)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    u.set_paragraph_spacing(p, after=0, line=1.08)
    r = p.add_run("Para uso compartido real y control de acceso, la publicación debe configurar almacenamiento Redis/KV, sesiones y credenciales de producción.")
    u.set_run_font(r, size=8.1, color=MUTED, italic=True)

    u.add_spacer(doc, 11)
    u.add_subheading(doc, "Rutina de continuidad", size=11.8, after=5)
    routine = doc.add_table(rows=1, cols=3)
    u.set_table_geometry(routine, [3246, 3247, 3247], indent=0)
    for idx, (title, body) in enumerate((
        ("CADA TURNO", "Actualizar stock, precio y visibilidad que hayan cambiado."),
        ("CADA SEMANA", "Revisar quiebres, solicitudes y consistencia entre locales."),
        ("ANTES DE CAMBIOS MASIVOS", "Descargar respaldo, validar el CSV y asignar un responsable."),
    )):
        cell = routine.cell(0, idx)
        u.shade_cell(cell, PALE_GREEN if idx == 1 else PALE)
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=110, bottom=110, start=120, end=120)
        p = u.clear_cell(cell)
        r = p.add_run(title + "\n")
        u.set_run_font(r, size=7.5, color=GREEN if idx == 1 else BLUE, bold=True)
        r = p.add_run(body)
        u.set_run_font(r, size=8.25, color=MUTED)


def page_decisions(doc):
    section_header(
        doc,
        "07",
        "Información útil para ordenar inventario y demanda",
        "El panel reúne señales operativas que ayudan a decidir qué revisar, dónde reponer y qué información publicada necesita ajustes.",
    )

    available = doc.add_table(rows=2, cols=4)
    u.set_table_geometry(available, [2435] * 4, indent=0)
    signals = (
        ("Valor referencial", "del inventario"),
        ("Unidades", "totales y por local"),
        ("Stock bajo", "y productos agotados"),
        ("Inventario", "por categoría"),
        ("Demanda", "de los últimos 14 días"),
        ("Más pedidos", "productos y unidades"),
        ("Solicitudes", "por sucursal"),
        ("Historial", "cantidad, unidades y valores"),
    )
    for idx, (label, value) in enumerate(signals):
        mini_stat(available.cell(idx // 4, idx % 4), label, value,
                  accent=GREEN if idx in (4, 5, 6) else BLUE,
                  fill=PALE_GREEN if idx in (4, 5, 6) else PALE)

    u.add_spacer(doc, 11)
    u.add_subheading(doc, "Cómo aprovechar esa información", size=12.1, after=6)
    cases = doc.add_table(rows=2, cols=2)
    u.set_table_geometry(cases, [4870, 4870], indent=0)
    items = (
        ("Priorizar reposición", "Revisar quiebres y productos con demanda registrada antes de planificar compras."),
        ("Orientar entre locales", "Detectar dónde existe disponibilidad y reforzar la coordinación entre sucursales."),
        ("Cuidar lo publicado", "Corregir precios, visibilidad o stock que no representen la situación del local."),
        ("Observar tendencias", "Comparar productos y sucursales para decidir qué información destacar o actualizar."),
    )
    for idx, (title, body) in enumerate(items):
        cell = cases.cell(idx // 2, idx % 2)
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=105, bottom=105, start=0 if idx % 2 == 0 else 175, end=120)
        p = u.clear_cell(cell)
        r = p.add_run(title + ". ")
        u.set_run_font(r, size=9.2, color=NAVY, bold=True)
        r = p.add_run(body)
        u.set_run_font(r, size=9.2, color=MUTED)
        add_bottom_rule(p, color=LINE, size=5, space=2)

    u.add_spacer(doc, 10)
    warning = doc.add_table(rows=1, cols=2)
    u.set_table_geometry(warning, [2050, 7690], indent=0)
    left, right = warning.cell(0, 0), warning.cell(0, 1)
    u.shade_cell(left, RED)
    u.remove_cell_borders(left)
    u.set_cell_margins(left, top=155, bottom=155, start=125, end=125)
    p = u.clear_cell(left)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DISTINCIÓN CLAVE")
    u.set_run_font(r, size=8.5, color=WHITE, bold=True)
    u.shade_cell(right, PALE_RED)
    u.remove_cell_borders(right)
    u.set_cell_margins(right, top=120, bottom=120, start=180, end=180)
    p = u.clear_cell(right)
    r = p.add_run("El historial registra solicitudes que llegaron al paso de WhatsApp. No equivale a ventas ni reservas confirmadas, no es un CRM y no guarda datos personales del cliente.")
    u.set_run_font(r, size=8.9, color=INK, bold=True)

    u.add_spacer(doc, 10)
    banner(doc, "Usar el panel para mejorar decisiones operativas, no para reemplazar la confirmación del equipo ni el inventario físico.", fill=NAVY, size=11.2)

    u.add_spacer(doc, 11)
    interpretation = doc.add_table(rows=1, cols=3)
    u.set_table_geometry(interpretation, [3246, 3247, 3247], indent=0)
    for idx, (title, body) in enumerate((
        ("SEÑAL", "Una solicitud muestra interés, no una venta cerrada."),
        ("CONTRASTE", "Comparar el historial con stock físico y respuesta del equipo."),
        ("ACCIÓN", "Corregir datos, reponer o reforzar la coordinación entre locales."),
    )):
        cell = interpretation.cell(0, idx)
        u.shade_cell(cell, LIGHT_BLUE if idx == 0 else (PALE_GREEN if idx == 2 else PALE))
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=110, bottom=110, start=120, end=120)
        p = u.clear_cell(cell)
        r = p.add_run(title + "\n")
        u.set_run_font(r, size=7.6, color=GREEN if idx == 2 else BLUE, bold=True)
        r = p.add_run(body)
        u.set_run_font(r, size=8.3, color=MUTED)


def page_activation(doc):
    section_header(
        doc,
        "08",
        "Publicar, llevarla al barrio y aprender de su uso",
        "La puesta en marcha combina preparación de datos, difusión local y revisión del primer mes. El objetivo es empezar simple y sostener la actualización.",
    )

    phases = doc.add_table(rows=1, cols=3)
    u.set_table_geometry(phases, [3246, 3247, 3247], indent=0)
    phase_items = (
        ("01", "PREPARAR", ["Validar identidad y datos.", "Cargar catálogo real.", "Configurar accesos y backend."], LIGHT_BLUE, BLUE),
        ("02", "ACTIVAR", ["QR en mesón, bolsas y comprobantes.", "Enlaces en WhatsApp, Google y redes.", "Capacitación breve por local."], PALE_GREEN, GREEN),
        ("03", "MEJORAR", ["Revisar los primeros 30 días.", "Corregir stock y catálogo.", "Ajustar la comunicación local."], WHITE, BLUE),
    )
    for idx, (num, title, lines, fill, accent) in enumerate(phase_items):
        block(phases.cell(0, idx), title, lines, fill=fill, accent=accent,
              numbered=num, body_size=8.65, border=True)

    u.add_spacer(doc, 10)
    metrics = doc.add_table(rows=1, cols=2)
    u.set_table_geometry(metrics, [4870, 4870], indent=0)
    block(metrics.cell(0, 0), "Disponible hoy en el panel", [
        "Stock, quiebres y alertas de reposición.",
        "Solicitudes registradas, productos y unidades.",
        "Distribución por sucursal y valores referenciales.",
    ], fill=LIGHT_BLUE, accent=BLUE, body_size=8.9)
    block(metrics.cell(0, 1), "Recomendado al publicar", [
        "Visitas, búsquedas, filtros y origen del tráfico.",
        "Clics a WhatsApp y tiempo de respuesta.",
        "Reservas confirmadas y recurrencia de clientes.",
    ], fill=PALE, accent=GREEN, body_size=8.9)

    u.add_spacer(doc, 8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    u.set_paragraph_spacing(p, after=8, line=1.05)
    r = p.add_run("Las métricas del segundo bloque requieren analítica o registro operativo adicional; no se presentan como funciones activas.")
    u.set_run_font(r, size=8.0, color=MUTED, italic=True)

    u.add_subheading(doc, "Ritmo recomendado del primer mes", size=12.1, after=6)
    rhythm = doc.add_table(rows=1, cols=4)
    u.set_table_geometry(rhythm, [2435] * 4, indent=0)
    for idx, (label, text) in enumerate((
        ("SEMANA 1", "Validar datos y capacitar"),
        ("SEMANA 2", "Instalar QR y publicar enlaces"),
        ("SEMANAS 3–4", "Difundir y observar consultas"),
        ("DÍA 30", "Revisar métricas y ajustar"),
    )):
        cell = rhythm.cell(0, idx)
        u.shade_cell(cell, NAVY if idx % 2 == 0 else BLUE)
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=125, bottom=125, start=90, end=90)
        p = u.clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label + "\n")
        u.set_run_font(r, size=7.7, color=WHITE, bold=True)
        r = p.add_run(text)
        u.set_run_font(r, size=7.9, color="EAF1F5")

    u.add_spacer(doc, 10)
    banner(doc, "Empezar simple, medir y mejorar con la experiencia real de los vecinos.", fill=BLUE, size=12.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    u.add_spacer(doc, 6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    u.set_paragraph_spacing(p, after=0, line=1.05)
    r = p.add_run("No se incluyen proyecciones de venta ni resultados garantizados.")
    u.set_run_font(r, size=8.0, color=MUTED, italic=True)

    u.add_spacer(doc, 11)
    u.add_subheading(doc, "Canales según el punto de contacto", size=11.8, after=5)
    channels = doc.add_table(rows=1, cols=3)
    u.set_table_geometry(channels, [3246, 3247, 3247], indent=0)
    for idx, (title, body) in enumerate((
        ("EN LA SUCURSAL", "QR en mesón, bolsas, comprobantes y carteles visibles."),
        ("EN CANALES PROPIOS", "WhatsApp Business, Google Business Profile, Instagram y Facebook."),
        ("EN LA COMUNIDAD", "Enlace en grupos vecinales y contenido útil para el sector cercano."),
    )):
        cell = channels.cell(0, idx)
        u.shade_cell(cell, PALE_GREEN if idx == 2 else PALE)
        u.remove_cell_borders(cell)
        u.set_cell_margins(cell, top=110, bottom=110, start=120, end=120)
        p = u.clear_cell(cell)
        r = p.add_run(title + "\n")
        u.set_run_font(r, size=7.5, color=GREEN if idx == 2 else BLUE, bold=True)
        r = p.add_run(body)
        u.set_run_font(r, size=8.25, color=MUTED)


def page_ready(doc):
    section_header(
        doc,
        "09",
        "Listo para avanzar con datos reales y una identidad confirmada",
        "La base funcional ya está construida. La etapa siguiente consiste en reemplazar datos demostrativos, configurar el entorno de producción y preparar a cada sucursal.",
    )

    status = doc.add_table(rows=1, cols=2)
    u.set_table_geometry(status, [4870, 4870], indent=0)
    block(status.cell(0, 0), "Construido y disponible", [
        "Catálogo responsive con búsqueda, sugerencias, categorías, filtros y ficha.",
        "Stock, precio y visibilidad por sucursal; pedido y WhatsApp.",
        "Panel, roles, CSV, respaldos, historial y SEO técnico base.",
        "Cuatro sucursales y 38 productos demo, incluidos 8 de Perfumería.",
    ], fill=PALE_GREEN, accent=GREEN, body_size=8.65)
    block(status.cell(0, 1), "Antes de producción", [
        "Recibir logo y colores oficiales de Farmacias Real.",
        "Cargar catálogo real y validar stock, precio, receta y bioequivalencia.",
        "Confirmar WhatsApp de Santa María y Simón Bolívar.",
        "Configurar dominio, almacenamiento compartido, sesiones y credenciales.",
        "Capacitar, publicar y revisar el primer mes.",
    ], fill=PALE_RED, accent=RED, body_size=8.65)

    u.add_spacer(doc, 10)
    u.add_subheading(doc, "Límites del alcance actual", size=12.1, after=6)
    limits = doc.add_table(rows=2, cols=3)
    u.set_table_geometry(limits, [3246, 3247, 3247], indent=0)
    labels = (
        "Sin pago online", "Sin despacho", "Sin diagnóstico o sustitución automática",
        "Sin integración POS/ERP", "Sin CRM", "Sin confirmación automática de WhatsApp",
    )
    for idx, label in enumerate(labels):
        cell = limits.cell(idx // 3, idx % 3)
        u.shade_cell(cell, PALE if idx % 2 == 0 else WHITE)
        u.set_cell_border(cell, color=LINE, size=6, sides=("bottom",))
        u.set_cell_margins(cell, top=110, bottom=110, start=110, end=110)
        p = u.clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        u.set_run_font(r, size=8.3, color=NAVY, bold=True)

    u.add_spacer(doc, 10)
    quality = doc.add_table(rows=1, cols=2)
    u.set_table_geometry(quality, [2450, 7290], indent=0)
    left, right = quality.cell(0, 0), quality.cell(0, 1)
    u.shade_cell(left, BLUE)
    u.remove_cell_borders(left)
    u.set_cell_margins(left, top=145, bottom=145, start=135, end=135)
    p = u.clear_cell(left)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("VALIDACIÓN ACTUAL")
    u.set_run_font(r, size=8.5, color=WHITE, bold=True)
    u.shade_cell(right, LIGHT_BLUE)
    u.remove_cell_borders(right)
    u.set_cell_margins(right, top=115, bottom=115, start=175, end=175)
    p = u.clear_cell(right)
    r = p.add_run("Build de producción y verificación de tipos aprobados; 4 pruebas puras de Perfumería aprobadas. La suite de 44 pruebas requiere una ejecución completa en un entorno que permita levantar el servidor.")
    u.set_run_font(r, size=8.6, color=INK, bold=True)

    u.add_spacer(doc, 10)
    banner(doc, "Validar datos  →  configurar  →  capacitar  →  publicar  →  medir", fill=NAVY, size=11.7, align=WD_ALIGN_PARAGRAPH.CENTER)

    u.add_spacer(doc, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    u.set_paragraph_spacing(p, after=5, line=1.1)
    r = p.add_run("Una experiencia digital simple que transforma la búsqueda en una conversación y la conversación en una visita a la farmacia.")
    u.set_run_font(r, size=13.2, color=NAVY, bold=True)
    p = doc.add_paragraph()
    add_logo_natural(p, width=1.89, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    u.set_paragraph_spacing(p, after=1, line=1.0)
    r = p.add_run("Recursos oficiales utilizados: logotipo horizontal e isotipo de Brida.")
    u.set_run_font(r, size=7.5, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    u.set_paragraph_spacing(p, after=0, line=1.0)
    r = p.add_run("Recurso de plataforma utilizado: imagen Open Graph del prototipo; la identidad de Farmacias Real incluida en ella es provisional.")
    u.set_run_font(r, size=7.5, color=MUTED, italic=True)


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_section(doc)

    props = doc.core_properties
    props.title = "Propuesta de plataforma digital para Farmacias Real"
    props.author = "Brida"
    props.subject = "Catálogo por sucursal, reserva por WhatsApp y control operativo"
    props.keywords = "Farmacias Real, Brida, catálogo, WhatsApp, sucursales, panel"
    props.comments = "Versión editorial profesional preparada por Brida."
    props.created = datetime(2026, 8, 18, 12, 0, 0)
    props.modified = datetime(2026, 8, 18, 12, 0, 0)

    pages = (
        page_cover,
        page_executive,
        page_journey,
        page_public_experience,
        page_branches,
        page_whatsapp,
        page_panel,
        page_decisions,
        page_activation,
        page_ready,
    )
    for idx, page in enumerate(pages):
        if idx:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.add_run().add_break(WD_BREAK.PAGE)
        page(doc)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_document()
