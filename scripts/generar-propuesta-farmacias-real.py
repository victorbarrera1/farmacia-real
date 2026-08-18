from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
OUT_DIR = REPO / "output" / "documentos"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_DOCX = OUT_DIR / "Propuesta_Plataforma_Digital_Farmacias_Real_Brida.docx"

BRIDA_LOGO = Path("/Users/victor/Downloads/WhatsApp Image 2026-08-09 at 12.37.29 PM-2.jpeg")
BRIDA_ICON = Path("/Users/victor/Downloads/WhatsApp Image 2026-08-09 at 12.37.29 PM.jpeg")
FARMACIA_OG = REPO / "public" / "og-farmacias-real.png"

# Preset: neighborhood_business_proposal (narrative_proposal alias).
# Named override: commercial_visual_compact_A4.
# The user explicitly requested a short, visual A4 sales proposal, so the A4
# geometry, left-aligned compact body rhythm, brand colors and card tables below
# intentionally override the Letter/long-prose defaults.
FONT = "Arial"
BRIDA_BLUE = "1769AA"
BRIDA_DARK = "111827"
PLATFORM_NAVY = "1B2A55"
PLATFORM_RED = "D6202A"
WA_GREEN = "147A50"
INK = "172033"
MUTED = "5E6B7A"
LIGHT_BLUE = "EAF4FB"
PALE = "F5F8FB"
PALE_GREEN = "EAF7F1"
PALE_RED = "FCEDEF"
LINE = "DDE5ED"
WHITE = "FFFFFF"

# A4 width 11906 DXA, 0.625 in side margins (900 DXA) => 10106 DXA usable.
TABLE_WIDTH = 9950
TABLE_INDENT = 120
CELL_MARGIN = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=None, color=INK, bold=None, italic=None, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.15, keep_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep_next:
        fmt.keep_with_next = True


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=100, start=CELL_MARGIN, bottom=100, end=CELL_MARGIN):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=LINE, size=8, sides=("top", "start", "bottom", "end")):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in sides:
        tag = qn(f"w:{side}")
        edge = borders.find(tag)
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def remove_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "start", "bottom", "end", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "nil")


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths, indent=TABLE_INDENT):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        set_row_cant_split(row)
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def clear_cell(cell):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p)
    return p


def add_text(doc, text, size=10.7, color=INK, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=5, line=1.15,
             keep_next=False):
    p = doc.add_paragraph()
    p.alignment = align
    set_paragraph_spacing(p, before, after, line, keep_next)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_kicker(doc, text, color=BRIDA_BLUE, align=WD_ALIGN_PARAGRAPH.LEFT, after=4):
    p = add_text(doc, text.upper(), size=8.6, color=color, bold=True, align=align, after=after, line=1.0, keep_next=True)
    p.paragraph_format.keep_with_next = True
    return p


def add_heading(doc, text, size=23, color=BRIDA_DARK, after=8):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=after, line=1.0, keep_next=True)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=True)
    return p


def add_subheading(doc, text, size=13, color=BRIDA_DARK, after=5):
    p = doc.add_paragraph(style="Heading 2")
    set_paragraph_spacing(p, before=4, after=after, line=1.05, keep_next=True)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=True)
    return p


def add_rich_paragraph(container, segments, size=10.5, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT,
                       before=0, after=4, line=1.15):
    p = container.add_paragraph() if hasattr(container, "add_paragraph") else container
    p.alignment = align
    set_paragraph_spacing(p, before, after, line)
    for text, bold, seg_color in segments:
        r = p.add_run(text)
        set_run_font(r, size=size, color=seg_color or color, bold=bold)
    return p


def add_picture_cropped(paragraph, path, width, height, crop, alt):
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width), height=Inches(height))
    blip_fill = shape._inline.xpath(".//pic:blipFill")[0]
    src_rect = OxmlElement("a:srcRect")
    for side, value in crop.items():
        src_rect.set(side, str(value))
    blip_fill.insert(1, src_rect)
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt)
    return shape


def add_logo(paragraph, width=3.7, height=1.12):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, after=4, line=1.0)
    return add_picture_cropped(
        paragraph,
        BRIDA_LOGO,
        width,
        height,
        {"l": 1200, "r": 1200, "t": 31500, "b": 31500},
        "Logotipo oficial de Brida",
    )


def add_icon(paragraph, width=0.23, height=0.23):
    return add_picture_cropped(
        paragraph,
        BRIDA_ICON,
        width,
        height,
        {"l": 24500, "r": 24500, "t": 24500, "b": 24500},
        "Isotipo oficial de Brida",
    )


def add_card(cell, number, title, body, fill=PALE, accent=BRIDA_BLUE, title_size=11.5, body_size=9.4):
    shade_cell(cell, fill)
    set_cell_border(cell, color=LINE, size=7)
    set_cell_margins(cell, top=150, bottom=150, start=170, end=170)
    clear_cell(cell)
    p_num = cell.paragraphs[0]
    set_paragraph_spacing(p_num, after=6, line=1.0)
    r = p_num.add_run(number)
    set_run_font(r, size=9, color=accent, bold=True)
    p_title = cell.add_paragraph()
    set_paragraph_spacing(p_title, after=4, line=1.0, keep_next=True)
    r = p_title.add_run(title)
    set_run_font(r, size=title_size, color=BRIDA_DARK, bold=True)
    p_body = cell.add_paragraph()
    set_paragraph_spacing(p_body, after=0, line=1.12)
    r = p_body.add_run(body)
    set_run_font(r, size=body_size, color=MUTED)


def add_two_col_callout(doc, left_title, left_lines, right_title, right_lines,
                        left_fill=LIGHT_BLUE, right_fill=PALE_GREEN):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [4975, 4975])
    for cell, title, lines, fill, accent in (
        (table.cell(0, 0), left_title, left_lines, left_fill, BRIDA_BLUE),
        (table.cell(0, 1), right_title, right_lines, right_fill, WA_GREEN),
    ):
        shade_cell(cell, fill)
        set_cell_border(cell, color=LINE, size=7)
        set_cell_margins(cell, top=170, bottom=170, start=190, end=190)
        clear_cell(cell)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=8, line=1.0, keep_next=True)
        r = p.add_run(title)
        set_run_font(r, size=12.2, color=accent, bold=True)
        for label, detail in lines:
            p = cell.add_paragraph()
            set_paragraph_spacing(p, after=5, line=1.12)
            r = p.add_run(label)
            set_run_font(r, size=9.6, color=BRIDA_DARK, bold=True)
            r = p.add_run(detail)
            set_run_font(r, size=9.6, color=MUTED)
    return table


def add_quote(doc, text, fill=PLATFORM_NAVY, color=WHITE, size=14.5):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    remove_cell_borders(cell)
    set_cell_margins(cell, top=220, bottom=220, start=260, end=260)
    p = clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0, line=1.1)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=True)
    return table


def add_spacer(doc, points=6):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=points, line=1.0)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_page_field(paragraph, field_name):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.7)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT
    h1._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h1.font.size = Pt(23)
    h1.font.bold = True
    h1.font.color.rgb = rgb(BRIDA_DARK)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT
    h2._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = rgb(BRIDA_DARK)
    h2.paragraph_format.space_before = Pt(4)
    h2.paragraph_format.space_after = Pt(5)
    h2.paragraph_format.line_spacing = 1.05
    h2.paragraph_format.keep_with_next = True

    # Real list styles are retained and normalized even though this visual
    # proposal mostly uses cards and labeled paragraphs instead of long lists.
    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(10.2)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def configure_section(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.2677)
    section.page_height = Inches(11.6929)
    section.top_margin = Inches(0.54)
    section.bottom_margin = Inches(0.50)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, after=2, line=1.0)
    r = p.add_run("FARMACIAS REAL  /  PROPUESTA DIGITAL")
    set_run_font(r, size=7.8, color=MUTED, bold=True)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), LINE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Inches(TABLE_WIDTH / 1440))
    set_table_geometry(table, [7100, 2850], indent=0)
    for cell in table.rows[0].cells:
        remove_cell_borders(cell)
        set_cell_margins(cell, top=20, bottom=20, start=0, end=0)
    left = clear_cell(table.cell(0, 0))
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_icon(left)
    r = left.add_run("  Brida · Propuesta comercial · 18 agosto 2026")
    set_run_font(r, size=8.2, color=MUTED)
    right = clear_cell(table.cell(0, 1))
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = right.add_run("Página ")
    set_run_font(r, size=8.2, color=MUTED)
    add_page_field(right, "PAGE")
    r = right.add_run(" de ")
    set_run_font(r, size=8.2, color=MUTED)
    add_page_field(right, "NUMPAGES")

    # Quiet first-page header, branded first-page footer.
    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    first_table = first_footer.add_table(rows=1, cols=2, width=Inches(TABLE_WIDTH / 1440))
    set_table_geometry(first_table, [7100, 2850], indent=0)
    for cell in first_table.rows[0].cells:
        remove_cell_borders(cell)
        set_cell_margins(cell, top=20, bottom=20, start=0, end=0)
    left = clear_cell(first_table.cell(0, 0))
    r = left.add_run("Brida · Propuesta preparada para Farmacias Real")
    set_run_font(r, size=8.2, color=MUTED)
    right = clear_cell(first_table.cell(0, 1))
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = right.add_run("Página ")
    set_run_font(r, size=8.2, color=MUTED)
    add_page_field(right, "PAGE")
    r = right.add_run(" de ")
    set_run_font(r, size=8.2, color=MUTED)
    add_page_field(right, "NUMPAGES")


def page_cover(doc):
    p = doc.add_paragraph()
    add_logo(p, width=3.55, height=1.02)
    add_spacer(doc, 8)
    add_kicker(doc, "Propuesta para Farmacias Real", align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

    hero = doc.add_table(rows=1, cols=1)
    set_table_geometry(hero, [TABLE_WIDTH])
    cell = hero.cell(0, 0)
    shade_cell(cell, PLATFORM_NAVY)
    remove_cell_borders(cell)
    set_cell_margins(cell, top=430, bottom=430, start=420, end=420)
    p = clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=12, line=1.0, keep_next=True)
    r = p.add_run("Farmacias Real,\nmás cerca de sus vecinos")
    set_run_font(r, size=27, color=WHITE, bold=True)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0, line=1.18)
    r = p.add_run("Catálogo por sucursal, reserva por WhatsApp y retiro presencial en una experiencia simple y cercana.")
    set_run_font(r, size=13, color="E7ECF7")

    add_spacer(doc, 13)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3316, 3317, 3317])
    items = (
        ("01", "ENCUENTRA", "Busca y filtra productos"),
        ("02", "RESERVA", "Envía el pedido por WhatsApp"),
        ("03", "RETIRA", "Confirma, paga y recibe en el local"),
    )
    for idx, (num, title, body) in enumerate(items):
        cell = table.cell(0, idx)
        add_card(cell, num, title, body, fill=LIGHT_BLUE if idx != 1 else PALE_RED,
                 accent=BRIDA_BLUE if idx != 1 else PLATFORM_RED, title_size=10.6, body_size=8.9)

    add_spacer(doc, 38)
    p = add_text(
        doc,
        "Una farmacia de barrio con la facilidad digital que hoy esperan los clientes.",
        size=15,
        color=BRIDA_DARK,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
        line=1.1,
    )
    p.paragraph_format.keep_with_next = True
    add_text(doc, "Presentado por Brida · Agosto 2026", size=9.4, color=MUTED,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_spacer(doc, 20)
    branches = doc.add_table(rows=1, cols=4)
    set_table_geometry(branches, [2487, 2488, 2487, 2488])
    for idx, label in enumerate(("Independencia 1443", "Sevilla 1201", "Santa María 1789", "Simón Bolívar 3751")):
        cell = branches.cell(0, idx)
        shade_cell(cell, PALE)
        set_cell_border(cell, color=LINE, size=6)
        set_cell_margins(cell, top=100, bottom=100, start=70, end=70)
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        set_run_font(r, size=8.2, color=MUTED, bold=True)


def page_opportunity(doc):
    add_kicker(doc, "01 · La oportunidad")
    add_heading(doc, "La decisión comienza antes de la visita")
    add_text(
        doc,
        "La plataforma responde tres preguntas antes de que el vecino salga de casa: qué necesita, en qué sucursal está disponible y cómo reservarlo.",
        size=12.2,
        color=MUTED,
        after=10,
        line=1.2,
    )

    frame = doc.add_table(rows=1, cols=1)
    set_table_geometry(frame, [TABLE_WIDTH])
    cell = frame.cell(0, 0)
    shade_cell(cell, PALE)
    set_cell_border(cell, color=LINE, size=8)
    set_cell_margins(cell, top=90, bottom=90, start=90, end=90)
    p = clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(FARMACIA_OG), width=Inches(6.78))
    shape._inline.docPr.set("descr", "Recurso visual real del prototipo de Farmacias Real")
    cap = cell.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(cap, before=3, after=0, line=1.0)
    r = cap.add_run("Recurso visual del prototipo funcional")
    set_run_font(r, size=8.2, color=MUTED, italic=True)

    add_spacer(doc, 10)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3316, 3317, 3317])
    cards = (
        ("MENOS TRASLADOS", "El cliente consulta disponibilidad antes de ir."),
        ("MEJORES CONSULTAS", "El personal recibe productos, cantidades y local elegidos."),
        ("MÁS OPORTUNIDADES", "Una búsqueda puede transformarse en una visita al local."),
    )
    for idx, (title, body) in enumerate(cards):
        add_card(table.cell(0, idx), f"0{idx+1}", title, body, fill=WHITE,
                 accent=PLATFORM_RED if idx == 2 else BRIDA_BLUE, title_size=9.8, body_size=8.8)
    add_spacer(doc, 9)
    add_quote(doc, "La web no reemplaza la farmacia física: la acerca.", fill=BRIDA_BLUE, size=15)


def page_journey(doc):
    add_kicker(doc, "02 · Experiencia del cliente")
    add_heading(doc, "Del buscador al mesón, sin pasos extra")
    add_text(doc, "Un recorrido breve, reconocible y pensado para funcionar desde el teléfono.",
             size=11.7, color=MUTED, after=10)

    table = doc.add_table(rows=2, cols=3)
    set_table_geometry(table, [3316, 3317, 3317])
    steps = (
        ("01", "Busca", "Por nombre, marca, presentación o principio activo."),
        ("02", "Elige sucursal", "Ve la información correspondiente al local de retiro."),
        ("03", "Revisa", "Consulta disponibilidad y precio referencial publicado."),
        ("04", "Prepara", "Agrega productos y ajusta las cantidades."),
        ("05", "Envía", "WhatsApp abre una solicitud ordenada y lista para revisar."),
        ("06", "Retira", "La farmacia confirma; el cliente paga y recibe en el local."),
    )
    for idx, item in enumerate(steps):
        add_card(table.cell(idx // 3, idx % 3), *item,
                 fill=LIGHT_BLUE if idx < 3 else WHITE,
                 accent=BRIDA_BLUE if idx != 4 else WA_GREEN,
                 title_size=11.5, body_size=9.0)

    add_spacer(doc, 11)
    add_quote(doc, "Sin cuentas, contraseñas, datos bancarios ni aplicaciones que instalar.",
              fill=PLATFORM_NAVY, size=14)
    add_spacer(doc, 9)
    add_two_col_callout(
        doc,
        "Conveniencia digital",
        [
            ("Búsqueda directa. ", "Resultados y filtros en la misma página."),
            ("Información previa. ", "Sucursal, disponibilidad y precio referencial."),
        ],
        "Operación presencial",
        [
            ("Confirmación humana. ", "La reserva se acuerda con la farmacia."),
            ("Control en el local. ", "La receta, el pago y la entrega se validan presencialmente."),
        ],
        left_fill=LIGHT_BLUE,
        right_fill=PALE_GREEN,
    )
    add_spacer(doc, 5)
    add_text(doc, "No es un e-commerce tradicional: la tecnología prepara la consulta y la farmacia conserva el control.",
             size=9.2, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_spacer(doc, 12)
    add_subheading(doc, "Diseñada para ser clara desde cualquier pantalla")
    clarity = doc.add_table(rows=1, cols=3)
    set_table_geometry(clarity, [3316, 3317, 3317])
    clarity_items = (
        ("A", "Móvil primero", "Búsqueda, sucursal y pedido permanecen a mano."),
        ("B", "Sin distracciones", "Sin carruseles automáticos ni movimientos inesperados."),
        ("C", "Uso amplio", "Textos legibles, objetivos táctiles y navegación por teclado."),
    )
    for idx, item in enumerate(clarity_items):
        add_card(clarity.cell(0, idx), *item, fill=WHITE if idx == 1 else PALE,
                 accent=BRIDA_BLUE, title_size=10.4, body_size=8.7)
    add_spacer(doc, 10)
    add_quote(doc, "Rapidez y claridad por sobre efectos innecesarios.", fill=BRIDA_BLUE, size=13.4)


def page_features(doc):
    add_kicker(doc, "03 · Propuesta de valor")
    add_heading(doc, "Todo lo necesario para decidir")
    add_text(doc, "La facilidad de una gran vitrina digital, manteniendo el trato directo de una farmacia de barrio.",
             size=11.7, color=MUTED, after=9)

    table = doc.add_table(rows=2, cols=3)
    set_table_geometry(table, [3316, 3317, 3317])
    features = (
        ("A", "Búsqueda rápida", "Nombre, marca, laboratorio, presentación o principio activo."),
        ("B", "Filtros útiles", "Disponibilidad, receta, precio, laboratorio y categorías."),
        ("C", "Bioequivalencia", "Muestra solo productos registrados con esa clasificación."),
        ("D", "Stock por local", "Cada sucursal publica su propia disponibilidad."),
        ("E", "Precio por sucursal", "El cliente revisa un valor referencial del local elegido."),
        ("F", "Alternativas de retiro", "Si falta stock, puede ver si existe en otra sucursal."),
    )
    for idx, item in enumerate(features):
        add_card(table.cell(idx // 3, idx % 3), *item,
                 fill=WHITE if idx % 2 else LIGHT_BLUE,
                 accent=PLATFORM_RED if idx == 2 else BRIDA_BLUE,
                 title_size=10.8, body_size=8.8)

    add_spacer(doc, 11)
    bio = doc.add_table(rows=1, cols=2)
    set_table_geometry(bio, [3170, 6780])
    left = bio.cell(0, 0)
    shade_cell(left, PLATFORM_RED)
    remove_cell_borders(left)
    set_cell_margins(left, top=200, bottom=200, start=210, end=210)
    p = clear_cell(left)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BIOEQUIVALENCIA")
    set_run_font(r, size=11.2, color=WHITE, bold=True)
    right = bio.cell(0, 1)
    shade_cell(right, PALE_RED)
    remove_cell_borders(right)
    set_cell_margins(right, top=180, bottom=180, start=220, end=220)
    p = clear_cell(right)
    r = p.add_run("Informa y ordena la búsqueda; no recomienda sustituciones ni reemplaza la orientación profesional.")
    set_run_font(r, size=10.2, color=INK, bold=True)

    add_spacer(doc, 10)
    add_subheading(doc, "Cuando el cliente cambia de sucursal, cambia la información que importa")
    table = doc.add_table(rows=1, cols=5)
    set_table_geometry(table, [1990, 1990, 1990, 1990, 1990])
    for idx, label in enumerate(("Catálogo", "Disponibilidad", "Precio", "Dirección", "WhatsApp")):
        cell = table.cell(0, idx)
        shade_cell(cell, PLATFORM_NAVY if idx % 2 == 0 else BRIDA_BLUE)
        remove_cell_borders(cell)
        set_cell_margins(cell, top=150, bottom=150, start=90, end=90)
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        set_run_font(r, size=9.2, color=WHITE, bold=True)
    add_spacer(doc, 12)
    add_two_col_callout(
        doc,
        "Lo que gana el vecino",
        [
            ("Encuentra más rápido. ", "Filtra antes de revisar todo el catálogo."),
            ("Decide mejor. ", "Consulta el local y la disponibilidad publicada."),
            ("Llega con certeza. ", "Puede coordinar antes de trasladarse."),
        ],
        "Lo que gana la farmacia",
        [
            ("Ordena consultas. ", "Recibe solicitudes más específicas."),
            ("Dirige la demanda. ", "Puede orientar hacia el local con stock."),
            ("Mantiene el control. ", "Confirma antes del pago y el retiro."),
        ],
        left_fill=LIGHT_BLUE,
        right_fill=PALE_GREEN,
    )
    add_spacer(doc, 10)
    add_quote(doc, "Tecnología para facilitar la atención, no para reemplazarla.", fill=PLATFORM_NAVY, size=13.4)


def page_whatsapp(doc):
    add_kicker(doc, "04 · Conversión")
    add_heading(doc, "WhatsApp transforma la búsqueda en conversación")
    add_text(
        doc,
        "El pedido llega preparado con productos, cantidades, sucursal, dirección y total referencial. La farmacia confirma disponibilidad, condiciones y valor final.",
        size=11.7,
        color=MUTED,
        after=10,
        line=1.18,
    )

    flow = doc.add_table(rows=1, cols=5)
    set_table_geometry(flow, [2800, 375, 2800, 375, 3600])
    flow_items = (
        (0, "CATÁLOGO", "El cliente prepara su pedido", PLATFORM_NAVY),
        (2, "WHATSAPP", "La solicitud llega ordenada", WA_GREEN),
        (4, "LOCAL", "La farmacia confirma y entrega", BRIDA_BLUE),
    )
    for idx, title, body, fill in flow_items:
        cell = flow.cell(0, idx)
        shade_cell(cell, fill)
        remove_cell_borders(cell)
        set_cell_margins(cell, top=180, bottom=180, start=150, end=150)
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title + "\n")
        set_run_font(r, size=10.3, color=WHITE, bold=True)
        r = p.add_run(body)
        set_run_font(r, size=8.6, color="EFF4F8")
    for idx in (1, 3):
        cell = flow.cell(0, idx)
        remove_cell_borders(cell)
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(">")
        set_run_font(r, size=16, color=MUTED, bold=True)

    add_spacer(doc, 12)
    add_two_col_callout(
        doc,
        "Para el cliente",
        [
            ("No reescribe. ", "Los productos y cantidades ya están incluidos."),
            ("Evita traslados a ciegas. ", "Recibe confirmación antes de ir."),
            ("Habla con una persona. ", "Usa un canal conocido y directo."),
        ],
        "Para la farmacia",
        [
            ("Recibe orden. ", "Conoce productos, cantidades y local de retiro."),
            ("Mantiene control. ", "Confirma stock, condición de venta y valor."),
            ("Sigue atendiendo. ", "El pago y la entrega permanecen en caja."),
        ],
    )

    add_spacer(doc, 10)
    add_quote(doc, "La tecnología prepara el pedido; la farmacia mantiene la relación con el cliente.",
              fill=WA_GREEN, size=13.8)
    add_spacer(doc, 9)
    safeguards = doc.add_table(rows=1, cols=3)
    set_table_geometry(safeguards, [3316, 3317, 3317])
    labels = (
        ("SIN PAGO ONLINE", "El valor final se confirma y se paga en el local."),
        ("SIN DESPACHO", "La entrega se realiza presencialmente en la sucursal."),
        ("RECETA PRESENCIAL", "La validación profesional sigue a cargo de la farmacia."),
    )
    for idx, (title, body) in enumerate(labels):
        add_card(safeguards.cell(0, idx), "", title, body, fill=PALE,
                 accent=BRIDA_BLUE, title_size=9.5, body_size=8.5)
    add_spacer(doc, 11)
    add_subheading(doc, "Cada solicitud llega con el contexto necesario")
    request_data = doc.add_table(rows=1, cols=5)
    set_table_geometry(request_data, [1990, 1990, 1990, 1990, 1990])
    for idx, label in enumerate(("Productos", "Cantidades", "Sucursal", "Dirección", "Total referencial")):
        cell = request_data.cell(0, idx)
        shade_cell(cell, PALE_GREEN if idx == 2 else WHITE)
        set_cell_border(cell, color=LINE, size=7)
        set_cell_margins(cell, top=135, bottom=135, start=70, end=70)
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        set_run_font(r, size=8.7, color=WA_GREEN if idx == 2 else BRIDA_DARK, bold=True)
    add_spacer(doc, 7)
    add_text(doc, "La solicitud no queda confirmada automáticamente: la farmacia responde antes de que el cliente se traslade.",
             size=8.8, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)


def page_admin(doc):
    add_kicker(doc, "05 · Operación")
    add_heading(doc, "Control simple para cuatro sucursales")
    add_text(doc, "El dueño conserva la visión general; cada encargado trabaja con el alcance correspondiente a su local.",
             size=11.7, color=MUTED, after=10)

    branches = doc.add_table(rows=1, cols=4)
    set_table_geometry(branches, [2487, 2488, 2487, 2488])
    branch_items = (
        ("Independencia 1443", "Independencia"),
        ("Sevilla 1201", "Independencia"),
        ("Santa María 1789", "Independencia"),
        ("Simón Bolívar 3751", "Ñuñoa"),
    )
    for idx, (name, comuna) in enumerate(branch_items):
        cell = branches.cell(0, idx)
        shade_cell(cell, PLATFORM_NAVY if idx % 2 == 0 else BRIDA_BLUE)
        remove_cell_borders(cell)
        set_cell_margins(cell, top=170, bottom=170, start=120, end=120)
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name + "\n")
        set_run_font(r, size=9.4, color=WHITE, bold=True)
        r = p.add_run(comuna)
        set_run_font(r, size=8.2, color="E7ECF7")

    add_spacer(doc, 12)
    cards = doc.add_table(rows=2, cols=2)
    set_table_geometry(cards, [4975, 4975])
    admin_features = (
        ("01", "Catálogo", "Crear, editar e importar productos por CSV; exportar cuando se necesite."),
        ("02", "Stock, precio y visibilidad", "Actualizar qué muestra cada local sin duplicar el catálogo."),
        ("03", "Roles por sucursal", "Administración general y encargados con permisos acotados."),
        ("04", "Reservas y respaldo", "Revisar solicitudes registradas y descargar una copia de los datos."),
    )
    for idx, item in enumerate(admin_features):
        add_card(cards.cell(idx // 2, idx % 2), *item,
                 fill=WHITE if idx in (1, 2) else LIGHT_BLUE,
                 accent=BRIDA_BLUE,
                 title_size=11.2, body_size=9.1)

    add_spacer(doc, 11)
    add_quote(doc, "Administrar la plataforma debe sentirse tan directo como actualizar una planilla, con controles que ayudan a evitar errores.",
              fill=BRIDA_BLUE, size=13.2)
    add_spacer(doc, 9)
    status = doc.add_table(rows=1, cols=2)
    set_table_geometry(status, [2600, 7350])
    left = status.cell(0, 0)
    shade_cell(left, PALE_GREEN)
    set_cell_border(left, color="B8DFCC", size=7)
    set_cell_margins(left, top=150, bottom=150, start=160, end=160)
    p = clear_cell(left)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PROTOTIPO FUNCIONAL")
    set_run_font(r, size=9.5, color=WA_GREEN, bold=True)
    right = status.cell(0, 1)
    shade_cell(right, PALE)
    set_cell_border(right, color=LINE, size=7)
    set_cell_margins(right, top=135, bottom=135, start=190, end=190)
    p = clear_cell(right)
    r = p.add_run("El modo compartido entre dispositivos requiere configurar almacenamiento y accesos antes de publicar.")
    set_run_font(r, size=9.4, color=MUTED)
    add_spacer(doc, 11)
    add_two_col_callout(
        doc,
        "Para el dueño",
        [
            ("Visión general. ", "Inventario, quiebres, stock bajo y demanda por sucursal."),
            ("Cambios controlados. ", "Productos, sucursales, importación y respaldos."),
        ],
        "Para cada local",
        [
            ("Alcance propio. ", "Stock, precio y visibilidad de su sucursal."),
            ("Solicitudes pertinentes. ", "Consulta las reservas asociadas a su local."),
        ],
        left_fill=LIGHT_BLUE,
        right_fill=PALE_GREEN,
    )


def page_activation(doc):
    add_kicker(doc, "06 · Puesta en marcha")
    add_heading(doc, "Cómo llevar la plataforma a los vecinos")
    add_text(doc, "Una activación pequeña, medible y fácil de sostener desde cada sucursal.",
             size=11.7, color=MUTED, after=10)

    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3316, 3317, 3317])
    phases = (
        ("01", "PREPARAR", "Validar identidad, contactos, horarios, catálogo, stock, precios y accesos."),
        ("02", "ACTIVAR", "QR en mesón y bolsas; enlace en WhatsApp Business, Google y redes sociales."),
        ("03", "MEJORAR", "Revisar los primeros 30 días y ajustar catálogo, stock y comunicación."),
    )
    for idx, item in enumerate(phases):
        add_card(table.cell(0, idx), *item,
                 fill=LIGHT_BLUE if idx == 0 else (PALE_GREEN if idx == 1 else WHITE),
                 accent=WA_GREEN if idx == 1 else BRIDA_BLUE,
                 title_size=11.2, body_size=9.0)

    add_spacer(doc, 13)
    add_subheading(doc, "Cinco señales para saber si está funcionando")
    kpis = doc.add_table(rows=1, cols=5)
    set_table_geometry(kpis, [1990, 1990, 1990, 1990, 1990])
    for idx, (label, detail) in enumerate((
        ("VISITAS", "al catálogo"),
        ("BÚSQUEDAS", "realizadas"),
        ("CLICS", "a WhatsApp"),
        ("RESERVAS", "confirmadas"),
        ("SIN STOCK", "más consultados"),
    )):
        cell = kpis.cell(0, idx)
        shade_cell(cell, PLATFORM_NAVY if idx % 2 == 0 else BRIDA_BLUE)
        remove_cell_borders(cell)
        set_cell_margins(cell, top=150, bottom=150, start=80, end=80)
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label + "\n")
        set_run_font(r, size=8.5, color=WHITE, bold=True)
        r = p.add_run(detail)
        set_run_font(r, size=7.8, color="E7ECF7")

    add_spacer(doc, 13)
    add_two_col_callout(
        doc,
        "Difusión local",
        [
            ("Dentro del local. ", "QR en mesón, bolsas y comprobantes."),
            ("Canales propios. ", "WhatsApp Business, Google Business Profile e Instagram."),
            ("Comunidad. ", "Información útil y publicaciones para el sector cercano."),
        ],
        "Decisiones con datos",
        [
            ("Stock. ", "Detectar productos consultados que aparecen agotados."),
            ("Contenido. ", "Reforzar categorías y búsquedas frecuentes."),
            ("Sucursal. ", "Identificar dónde hace falta más difusión o respuesta."),
        ],
        left_fill=LIGHT_BLUE,
        right_fill=PALE_GREEN,
    )
    add_spacer(doc, 10)
    add_quote(doc, "Empezar simple, medir y mejorar con la experiencia real de los vecinos.",
              fill=PLATFORM_NAVY, size=13.8)
    add_spacer(doc, 6)
    add_text(doc, "No se incluyen proyecciones de ventas ni resultados garantizados; el impacto debe medirse después del lanzamiento.",
             size=8.8, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_spacer(doc, 11)
    add_subheading(doc, "Ritmo recomendado del primer mes")
    rhythm = doc.add_table(rows=1, cols=4)
    set_table_geometry(rhythm, [2487, 2488, 2487, 2488])
    rhythm_items = (
        ("SEMANA 1", "Validar datos y capacitar"),
        ("SEMANA 2", "Instalar QR y publicar enlaces"),
        ("SEMANAS 3-4", "Difundir y observar consultas"),
        ("DÍA 30", "Revisar métricas y ajustar"),
    )
    for idx, (label, body) in enumerate(rhythm_items):
        cell = rhythm.cell(0, idx)
        shade_cell(cell, PLATFORM_NAVY if idx % 2 == 0 else BRIDA_BLUE)
        remove_cell_borders(cell)
        set_cell_margins(cell, top=140, bottom=140, start=95, end=95)
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label + "\n")
        set_run_font(r, size=8.6, color=WHITE, bold=True)
        r = p.add_run(body)
        set_run_font(r, size=8.1, color="E7ECF7")


def page_next_steps(doc):
    add_kicker(doc, "07 · Próximo paso")
    add_heading(doc, "Publicar con datos reales y una identidad confirmada")
    add_text(
        doc,
        "Farmacias Real puede competir en facilidad digital sin perder la confianza, la cercanía y la atención personal de una farmacia de barrio.",
        size=12.2,
        color=MUTED,
        after=10,
        line=1.18,
    )

    add_two_col_callout(
        doc,
        "Lo que ya está construido",
        [
            ("Experiencia pública. ", "Búsqueda, filtros, sucursales, pedido y WhatsApp."),
            ("Operación. ", "Panel, roles, stock, precios, visibilidad, CSV y respaldos."),
            ("Base de publicación. ", "SEO técnico, diseño responsive y build de producción validado."),
        ],
        "Lo que debe confirmar el cliente",
        [
            ("Identidad. ", "Logo y colores oficiales de Farmacias Real."),
            ("Datos. ", "Catálogo, stock, precios, horarios y condiciones de venta."),
            ("Contactos. ", "WhatsApp definitivo de Santa María y Simón Bolívar."),
            ("Publicación. ", "Dominio, almacenamiento compartido y accesos del panel."),
        ],
        left_fill=PALE_GREEN,
        right_fill=PALE_RED,
    )

    add_spacer(doc, 11)
    facts = doc.add_table(rows=2, cols=2)
    set_table_geometry(facts, [2500, 7450])
    rows = (
        ("DATOS COMPROBADOS", "El repositorio contiene cuatro sucursales y un catálogo demostrativo. El build de producción y el chequeo de tipos pasaron el 18 de agosto de 2026."),
        ("ESCENARIOS", "Los beneficios comerciales y la adopción esperada son objetivos a medir, no cifras garantizadas. Esta propuesta no incluye proyecciones de ventas o conversión."),
    )
    for idx, (label, body) in enumerate(rows):
        left = facts.cell(idx, 0)
        right = facts.cell(idx, 1)
        shade_cell(left, PLATFORM_NAVY if idx == 0 else BRIDA_BLUE)
        shade_cell(right, PALE)
        remove_cell_borders(left)
        remove_cell_borders(right)
        set_cell_margins(left, top=145, bottom=145, start=160, end=160)
        set_cell_margins(right, top=145, bottom=145, start=190, end=190)
        p = clear_cell(left)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        set_run_font(r, size=8.8, color=WHITE, bold=True)
        p = clear_cell(right)
        r = p.add_run(body)
        set_run_font(r, size=9.2, color=MUTED)

    add_spacer(doc, 10)
    add_quote(doc, "Una experiencia digital simple que transforma la búsqueda en una conversación y la conversación en una visita a la farmacia.",
              fill=BRIDA_BLUE, size=13.4)
    add_spacer(doc, 9)
    p = doc.add_paragraph()
    add_logo(p, width=2.45, height=0.72)
    add_text(doc, "Recursos utilizados: logotipo horizontal e isotipo oficial de Brida; recurso Open Graph y datos funcionales del prototipo Farmacias Real.",
             size=8.2, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_text(doc, "Nota: el emblema actual de Farmacias Real es provisional y no se presenta como identidad oficial del cliente.",
             size=8.2, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_section(doc)

    props = doc.core_properties
    props.title = "Propuesta de plataforma digital para Farmacias Real"
    props.author = "Brida"
    props.subject = "Catálogo por sucursal, reserva por WhatsApp y retiro presencial"
    props.keywords = "Farmacias Real, Brida, propuesta digital, WhatsApp, catálogo, sucursales"
    props.comments = "Propuesta comercial breve y editable preparada por Brida."
    props.created = datetime(2026, 8, 18, 12, 0, 0)
    props.modified = datetime(2026, 8, 18, 12, 0, 0)

    page_cover(doc)
    add_page_break(doc)
    page_opportunity(doc)
    add_page_break(doc)
    page_journey(doc)
    add_page_break(doc)
    page_features(doc)
    add_page_break(doc)
    page_whatsapp(doc)
    add_page_break(doc)
    page_admin(doc)
    add_page_break(doc)
    page_activation(doc)
    add_page_break(doc)
    page_next_steps(doc)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_document()
